"""
Document ingestion module for TXT and DOCX files
Handles text extraction, chunking, and vectorization
"""
import numpy as np
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from datetime import datetime
import hashlib
import re

# Document readers
import docx  # python-docx

from src.utils.gemini_client import gemini_client
from src.utils.models import DocumentMetadata, DocumentChunkMetadata
from src.vectordb.vector_store import vector_store_manager


class DocumentIngestion:
    """Handles document file ingestion and vectorization"""
    
    def __init__(self):
        """Initialize ingestion handler"""
        self.documents: Dict[str, str] = {}  # file_id -> full text
        self.document_metadata: Dict[str, DocumentMetadata] = {}
        self.document_chunks: Dict[str, List[str]] = {}  # file_id -> list of chunks
    
    def generate_file_id(self, filename: str) -> str:
        """
        Generate unique file ID from filename
        
        Args:
            filename: Original filename
            
        Returns:
            Unique file identifier
        """
        unique_str = f"{filename}_{datetime.now().isoformat()}"
        file_hash = hashlib.md5(unique_str.encode()).hexdigest()[:8]
        clean_name = Path(filename).stem.replace(" ", "_")
        return f"doc_{clean_name}_{file_hash}"
    
    def read_txt(self, filepath: str) -> str:
        """
        Read text from TXT file
        
        Args:
            filepath: Path to TXT file
            
        Returns:
            Extracted text
        """
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def read_docx(self, filepath: str) -> str:
        """
        Read text from DOCX file
        
        Args:
            filepath: Path to DOCX file
            
        Returns:
            Extracted text
        """
        doc = docx.Document(filepath)
        
        # Extract all paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)
        
        # Combine all text
        full_text = "\n".join(paragraphs)
        if table_texts:
            full_text += "\n\nTables:\n" + "\n".join(table_texts)
        
        return full_text
    
    def chunk_text(
        self, 
        text: str, 
        chunk_size: int = 1000, 
        overlap: int = 200
    ) -> List[str]:
        """
        Split text into overlapping chunks
        
        Args:
            text: Input text
            chunk_size: Target chunk size in characters
            overlap: Overlap between chunks
            
        Returns:
            List of text chunks
        """
        # Split by paragraphs first
        paragraphs = text.split('\n')
        
        chunks = []
        current_chunk = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # If adding this paragraph exceeds chunk size, save current chunk
            if len(current_chunk) + len(para) > chunk_size and current_chunk:
                chunks.append(current_chunk)
                
                # Start new chunk with overlap
                words = current_chunk.split()
                overlap_text = " ".join(words[-overlap//5:])  # Approximate word overlap
                current_chunk = overlap_text + " " + para
            else:
                current_chunk += "\n" + para if current_chunk else para
        
        # Add final chunk
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    def extract_questions_and_answers(self, text: str) -> List[Dict[str, str]]:
        """
        Extract Q&A pairs from document (for your prompt document format)
        
        Args:
            text: Document text
            
        Returns:
            List of {question, answer, analysis} dictionaries
        """
        qa_pairs = []
        
        # Pattern to match Q1:, Q2:, etc.
        question_pattern = r'Q\d+:\s*(.+?)(?=\n\nAns:|$)'
        answer_pattern = r'Ans:\s*(.+?)(?=\n\nANALYSIS:|Q\d+:|$)'
        analysis_pattern = r'ANALYSIS:\s*(.+?)(?=Q\d+:|$)'
        
        # Find all questions
        questions = re.finditer(question_pattern, text, re.DOTALL)
        
        for q_match in questions:
            question_text = q_match.group(1).strip()
            start_pos = q_match.start()
            
            # Find corresponding answer
            answer_match = re.search(answer_pattern, text[start_pos:], re.DOTALL)
            answer_text = answer_match.group(1).strip() if answer_match else ""
            
            # Find corresponding analysis
            analysis_match = re.search(analysis_pattern, text[start_pos:], re.DOTALL)
            analysis_text = analysis_match.group(1).strip() if analysis_match else ""
            
            qa_pairs.append({
                'question': question_text,
                'answer': answer_text,
                'analysis': analysis_text
            })
        
        return qa_pairs
    
    def ingest_document(
        self, 
        filepath: str, 
        file_id: Optional[str] = None,
        vectorize: bool = True,
        chunk_size: int = 1000
    ) -> Tuple[str, DocumentMetadata]:
        """
        Ingest document file and optionally vectorize
        
        Args:
            filepath: Path to document file
            file_id: Optional custom file ID
            vectorize: Whether to create embeddings
            chunk_size: Size of text chunks for embedding
            
        Returns:
            Tuple of (file_id, metadata)
        """
        filepath = Path(filepath)
        
        # Read document based on extension
        if filepath.suffix.lower() == '.txt':
            text = self.read_txt(str(filepath))
            doc_type = 'txt'
        elif filepath.suffix.lower() == '.docx':
            text = self.read_docx(str(filepath))
            doc_type = 'docx'
        else:
            raise ValueError(f"Unsupported file type: {filepath.suffix}")
        
        # Generate file ID if not provided
        if file_id is None:
            file_id = self.generate_file_id(filepath.name)
        
        # Extract Q&A pairs if present
        qa_pairs = self.extract_questions_and_answers(text)
        
        # Create chunks
        chunks = self.chunk_text(text, chunk_size=chunk_size)
        
        # Create metadata
        metadata = DocumentMetadata(
            file_id=file_id,
            filename=filepath.name,
            document_type=doc_type,
            num_characters=len(text),
            num_chunks=len(chunks),
            num_qa_pairs=len(qa_pairs),
            ingestion_timestamp=datetime.now().isoformat(),
            has_questions=len(qa_pairs) > 0
        )
        
        # Store document and metadata
        self.documents[file_id] = text
        self.document_metadata[file_id] = metadata
        self.document_chunks[file_id] = chunks
        
        print(f"✓ Loaded {filepath.name}: {len(text)} chars, {len(chunks)} chunks, {len(qa_pairs)} Q&A pairs")
        
        # Vectorize if requested
        if vectorize:
            self._vectorize_document(text, chunks, qa_pairs, file_id, metadata)
        
        return file_id, metadata
    
    def _vectorize_document(
        self, 
        full_text: str,
        chunks: List[str],
        qa_pairs: List[Dict[str, str]],
        file_id: str,
        metadata: DocumentMetadata
    ) -> None:
        """
        Create embeddings for document chunks and Q&A pairs
        
        Args:
            full_text: Full document text
            chunks: Text chunks
            qa_pairs: Extracted Q&A pairs
            file_id: File identifier
            metadata: Document metadata
        """
        print(f"Creating embeddings for {file_id}...")
        
        # Create vector store
        store = vector_store_manager.create_store(file_id)
        
        vectors_list = []
        metadata_list = []
        
        # 1. Vectorize text chunks
        print(f"  - Vectorizing {len(chunks)} chunks...")
        for idx, chunk in enumerate(chunks):
            embedding = gemini_client.generate_embedding(chunk)
            vectors_list.append(embedding)
            
            vec_meta = DocumentChunkMetadata(
                file_id=file_id,
                chunk_idx=idx,
                chunk_type='text',
                original_text=chunk[:500],
                question_text=None,
                answer_text=None
            )
            metadata_list.append(vec_meta)
            
            if (idx + 1) % 10 == 0:
                print(f"    Processed {idx + 1}/{len(chunks)} chunks")
        
        # 2. Vectorize Q&A pairs separately for better retrieval
        if qa_pairs:
            print(f"  - Vectorizing {len(qa_pairs)} Q&A pairs...")
            for idx, qa in enumerate(qa_pairs):
                # Create combined text for embedding
                qa_text = f"Question: {qa['question']}\nAnswer: {qa['answer']}"
                if qa['analysis']:
                    qa_text += f"\nAnalysis: {qa['analysis']}"
                
                embedding = gemini_client.generate_embedding(qa_text)
                vectors_list.append(embedding)
                
                vec_meta = DocumentChunkMetadata(
                    file_id=file_id,
                    chunk_idx=len(chunks) + idx,
                    chunk_type='qa_pair',
                    original_text=qa_text[:500],
                    question_text=qa['question'],
                    answer_text=qa['answer'],
                    analysis_text=qa.get('analysis', '')
                )
                metadata_list.append(vec_meta)
        
        # Add all vectors to store
        vectors_array = np.vstack(vectors_list)
        store.add_vectors(vectors_array, metadata_list)
        
        # Save to disk
        vector_store_manager.save_store(file_id)
        
        print(f"✓ Created {len(vectors_list)} embeddings for {file_id}")
    
    def get_document_text(self, file_id: str) -> str:
        """Get full document text by file_id"""
        if file_id not in self.documents:
            raise ValueError(f"Document {file_id} not loaded")
        return self.documents[file_id]
    
    def get_document_chunks(self, file_id: str) -> List[str]:
        """Get document chunks by file_id"""
        if file_id not in self.document_chunks:
            raise ValueError(f"Document {file_id} not loaded")
        return self.document_chunks[file_id]
    
    def get_metadata(self, file_id: str) -> DocumentMetadata:
        """Get metadata by file_id"""
        if file_id not in self.document_metadata:
            raise ValueError(f"Metadata for {file_id} not found")
        return self.document_metadata[file_id]
    
    def search_document(
        self, 
        query: str, 
        file_id: Optional[str] = None,
        top_k: int = 3
    ) -> List[Tuple[str, float, Dict]]:
        """
        Search for relevant chunks in documents
        
        Args:
            query: Search query
            file_id: Optional file to search in
            top_k: Number of results
            
        Returns:
            List of (chunk_text, similarity_score, metadata) tuples
        """
        # Generate query embedding
        query_vector = gemini_client.generate_query_embedding(query)
        
        results = []
        
        # Search in specific file or all files
        file_ids = [file_id] if file_id else list(self.documents.keys())
        
        for fid in file_ids:
            store = vector_store_manager.get_store(fid)
            if not store:
                continue
            
            # Search
            search_results = store.search(query_vector, k=top_k, file_id=fid)
            
            for meta, distance in search_results:
                similarity = 1 - distance
                results.append((
                    meta.original_text,
                    float(similarity),
                    {
                        'file_id': meta.file_id,
                        'chunk_type': meta.chunk_type,
                        'question': getattr(meta, 'question_text', None),
                        'answer': getattr(meta, 'answer_text', None),
                        'analysis': getattr(meta, 'analysis_text', None)
                    }
                ))
        
        # Sort by similarity
        results.sort(key=lambda x: x[1], reverse=True)
        
        return results[:top_k]
    
    def list_documents(self) -> List[Dict[str, any]]:
        """List all loaded documents"""
        return [
            {
                "file_id": fid,
                "filename": meta.filename,
                "type": meta.document_type,
                "characters": meta.num_characters,
                "chunks": meta.num_chunks,
                "qa_pairs": meta.num_qa_pairs
            }
            for fid, meta in self.document_metadata.items()
        ]


# Global document ingestion instance
document_ingestion = DocumentIngestion()